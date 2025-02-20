
# from docx import Document
# from keybert import KeyBERT

# # Step 1: Read text from the .docx file
# file_path = "/Users/khyati/Downloads/example.docx"
# doc = Document(file_path)
# text = " ".join([para.text for para in doc.paragraphs if para.text.strip()])  # Extracted text

# # Step 2: Load BERT-based keyword extractor
# kw_model = KeyBERT()

# # Step 3: Extract keywords (top 5)
# bert_keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words="english", top_n=3)
# bert_keywords = [kw[0] for kw in bert_keywords]

# # Step 4: Print results
# print("BERT Keywords:", bert_keywords)

# -----------------------------------------------
from docx import Document
from keybert import KeyBERT
import pandas as pd
import re
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer


# Step 1: Read text from the .docx file
file_path = "/Users/khyati/Documents/Resumes/DE_Zoomcamp/sample-work-DE/KeyWordExtraction/example_changed_test.docx"
doc = Document(file_path)
text = " ".join([para.text for para in doc.paragraphs if para.text.strip()])  # Extracted text



# Step 2: Load the exception list from Excel
exception_file = "/Users/khyati/Documents/Resumes/DE_Zoomcamp/sample-work-DE/KeyWordExtraction/Exception list Food and Beverages New.xlsx"
exception_df = pd.read_excel(exception_file, usecols=[2, 4])  # Read only 3rd and 5th columns

# Clean the exception list by removing numbered prefixes (e.g., '1. Keyword') and converting to lowercase
# exception_list = set(
#     exception_df.stack()
#     .dropna()
#     .astype(str)
#     .apply(lambda x: re.sub(r'^\d+\.\s*', '', x).strip().lower())  # Remove numbered prefixes like '1. '
# )
# print("Cleaned Exception List:", exception_list)


exception_df = pd.read_excel(exception_file, usecols=[2, 4])  # Read only 3rd and 5th columns

# Clean the exception list by removing numbered prefixes (e.g., '1. Keyword') and converting to lowercase
exception_list = set(
    exception_df.stack()
    .dropna()
    .astype(str)
    .apply(lambda x: re.sub(r'^\d+\.\s*', '', x).strip().lower())  # Remove numbered prefixes like '1. '
)

# Step 3: Extract keywords using KeyBERT
kw_model = KeyBERT()
bert_keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words="english", top_n=20)
bert_keywords = [kw[0].lower() for kw in bert_keywords]

# Step 4: Properly clean and extract words from the example text
# Here, we'll treat multi-word names as valid words, so we will use a regular expression for that.
example_text = " ".join([para.text for para in doc.paragraphs if para.text.strip()])  # Get all text
example_words = re.findall(r'\b[\w\'-]+\b', example_text.lower())  # This will match words including hyphenated ones like "Yellow-Fin"
example_word_counts = Counter(example_words)

# Debug: Print cleaned example words and their counts
print("Example Word Counts:", example_word_counts)

# Remove words that appear too frequently in the example text
threshold = 5  # We can adjust this threshold depending on how often the terms like "Bagels" and "Sashimi" appear
frequent_words = {word for word, count in example_word_counts.items() if count >= threshold}

# Step 5: Compare and remove any matching words from the extracted keywords and frequent words
valid_keywords = [
    kw for kw in bert_keywords if kw not in exception_list and kw not in frequent_words
]

# Step 6: Print final extracted valid keywords
print("Valid Keywords:", valid_keywords)

