
# # Approach : 
# Step-by-Step Breakdown
# 1. Load and Read the Document

# from docx import Document

# file_path = "/Users/khyati/Downloads/Example one record product description for Summarization.docx"
# doc = Document(file_path)
# text = " ".join([para.text for para in doc.paragraphs if para.text.strip()])
# We loaded the .docx file and extracted text from all paragraphs.
# The text variable now contains the entire document as a string.
# 2. Apply TF-IDF Vectorization

# from sklearn.feature_extraction.text import TfidfVectorizer

# vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
# X = vectorizer.fit_transform([text])
# TF-IDF (Term Frequency - Inverse Document Frequency) is used to find important words in the text.
# TfidfVectorizer converts text into a numerical representation based on:
# Term Frequency (TF): How often a word appears in the text.
# Inverse Document Frequency (IDF): Gives lower scores to common words and higher scores to unique words.
# Parameters used:
# stop_words="english" → Removes common words like "the," "is," etc.
# ngram_range=(1,2) → Extracts both single words (unigrams) and two-word phrases (bigrams).
# The X variable now holds a TF-IDF matrix, where each word/phrase has a weight.
# 3. Extract the Top Keywords

# tfidf_scores = dict(zip(vectorizer.get_feature_names_out(), X.toarray()[0]))
# tfidf_keywords = sorted(tfidf_scores, key=tfidf_scores.get, reverse=True)[:5]
# get_feature_names_out() retrieves all words/phrases analyzed by TF-IDF.
# X.toarray()[0] converts the sparse TF-IDF matrix into an array of scores for each word.
# zip() combines words with their TF-IDF scores.
# sorted(..., reverse=True)[:5] sorts the words by highest TF-IDF score and selects the top 5 most relevant keywords.
# Final Output:
# The tfidf_keywords list contains 1-2 word phrases that are most important in describing the product.
# These keywords summarize the key aspects of the text.
# Example Output
# If the document contains:

# "DCSL Double Distilled Arrack, DCSL Very Special Old Arrack, DCSL Extra Special Arrack, Elephant House Necto, Frozen Thalapath, Keells Chicken Meatballs"
# The TF-IDF extracted keywords might be:

# ["DCSL Arrack", "Frozen Thalapath", "Keells Chicken", "Special Arrack", "Elephant House"]
# These highlight key product names while ignoring common words.
# Your extracted keywords are:
# ['finagle', 'roti', 'dcsl']
# This suggests that TF-IDF has identified words that appear frequently and uniquely in your document. Let’s analyze why this happened and how we can improve it.

# Why Did These Words Appear?
# "Finagle" – Likely a brand name that appears multiple times.
# "Roti" – A common product type that is repeated often in the document.
# "DCSL" – A company or brand name that appears uniquely in the text.
# TF-IDF focuses on words that appear frequently but are not generic stopwords. This is why it picked brand and product names.

# How to Improve Keyword Extraction?
# If you want more descriptive keywords (e.g., "DCSL Arrack" instead of just "DCSL"), try these improvements:

# 1. Increase N-gram Range

# Instead of using single words, extract two-word or three-word phrases:

# vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(2, 3))
# This will allow phrases like:

# "DCSL Arrack" instead of just "DCSL"
# "Finagle Roti" instead of "Finagle"
# 2. Remove Less Informative Words

# Some words (like "roti") might not be useful on their own. You can manually filter out common words after extraction:

# common_words = {"roti", "beer", "tea"}  # Add words that don’t add value
# filtered_keywords = [kw for kw in tfidf_keywords if kw.lower() not in common_words]
# This will keep brand names and unique terms while removing generic words.

# 3. Combine with Named Entity Recognition (NER)

# Using SpaCy, we can extract brands and product names explicitly:

# import spacy
# nlp = spacy.load("en_core_web_sm")
# doc_nlp = nlp(text)
# spacy_keywords = [ent.text for ent in doc_nlp.ents if ent.label_ in ["PRODUCT", "ORG"]]
# This will give words like "DCSL Arrack" or "Elephant House" instead of single words.

# Summary
# To get better keywords, modify your approach: ✔ Use bigrams & trigrams (ngram_range=(2,3))
# ✔ Filter out generic words manually
# ✔ Use Named Entity Recognition (NER) with SpaCy

# --------------------------------------------------------------


from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer

file_path = file_path = "/Users/khyati/Downloads/example.docx"

doc = Document(file_path)
text = " ".join([para.text for para in doc.paragraphs if para.text.strip()])

# vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(2, 3))

X = vectorizer.fit_transform([text])

tfidf_scores = dict(zip(vectorizer.get_feature_names_out(), X.toarray()[0]))
tfidf_keywords = sorted(tfidf_scores, key=tfidf_scores.get, reverse=True)[:3]
print("Extracted Keywords:", tfidf_keywords)

# ---------------------------------------------------------------
# from sklearn.feature_extraction.text import TfidfVectorizer


# # Step 1: Read text from the .docx file
# file_path = "/Users/khyati/Documents/Resumes/DE_Zoomcamp/sample-work-DE/KeyWordExtraction/example_changed_test.docx"
# doc = Document(file_path)
# text = " ".join([para.text for para in doc.paragraphs if para.text.strip()])  # Extracted text



# # Step 2: Load the exception list from Excel
# exception_file = "/Users/khyati/Documents/Resumes/DE_Zoomcamp/sample-work-DE/KeyWordExtraction/Exception list Food and Beverages New.xlsx"
# exception_df = pd.read_excel(exception_file, usecols=[2, 4])  # Read only 3rd and 5th columns

# # Clean the exception list by removing numbered prefixes (e.g., '1. Keyword') and converting to lowercase
# # exception_list = set(
# #     exception_df.stack()
# #     .dropna()
# #     .astype(str)
# #     .apply(lambda x: re.sub(r'^\d+\.\s*', '', x).strip().lower())  # Remove numbered prefixes like '1. '
# # )
# # print("Cleaned Exception List:", exception_list)


# exception_df = pd.read_excel(exception_file, usecols=[2, 4])
# exception_df = pd.read_excel(exception_file, usecols=[2, 4])  # Read only 3rd and 5th columns

# # Clean the exception list by removing numbered prefixes (e.g., '1. Keyword') and converting to lowercase
# exception_list = set(
#     exception_df.stack()
#     .dropna()
#     .astype(str)
#     .apply(lambda x: re.sub(r'^\d+\.\s*', '', x).strip().lower())  # Remove numbered prefixes like '1. '
# )

# vectorizer = TfidfVectorizer(ngram_range=(1, 2), stop_words="english")
# X = vectorizer.fit_transform([text])

# # Get feature names (words/phrases)
# features = vectorizer.get_feature_names_out()

# # Get TF-IDF scores for each feature
# tfidf_scores = X.toarray().flatten()

# # Create a DataFrame to hold the keywords and their corresponding TF-IDF scores
# keywords_df = pd.DataFrame({
#     'keyword': features,
#     'score': tfidf_scores
# })

# # Step 4: Sort keywords by their TF-IDF score in descending order
# keywords_df = keywords_df.sort_values(by='score', ascending=False)

# # Step 5: Remove keywords that are in the exception list
# valid_keywords = keywords_df[~keywords_df['keyword'].str.lower().isin(exception_list)]

# # Step 6: Get the top valid keywords
# top_valid_keywords = valid_keywords.head(10)  # Adjust the number of top keywords as needed

# # Step 7: Display the results
# print("Valid Keywords Extracted Using TF-IDF:")
# print(top_valid_keywords)
